from django.shortcuts import render, render_to_response, RequestContext
from django.contrib import messages
from django.core.mail import send_mail
from django.conf import settings
from django.http import HttpResponse
from django.contrib import auth
from django.shortcuts import render,render_to_response,redirect,get_object_or_404
from django.contrib.auth import  authenticate
from django.contrib.auth.models import User
from django.contrib.auth.decorators import login_required,user_passes_test,permission_required
from django.core.urlresolvers import reverse
# from salidas.forms import *
from salidas.models import *
from datetime import date
import locale
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from io import StringIO
from docx import * #to generate Docs
import io,os,os.path,tempfile, zipfile
import traceback

locale.setlocale(locale.LC_ALL, '')

def datetimeToString(datetime):
    monthDict={1:'Enero', 2:'Febrero', 3:'Marzo', 4:'Abril', 5:'Mayo', 6:'Junio', 7:'Julio', 8:'Agosto', 9:'Septiembre', 10:'Octubre', 11:'Noviembre', 12:'Diciembre'}
    day = int(datetime.day)
    month = monthDict[int(datetime.month)]
    year = int(datetime.year)
    return str(day)+" de "+str(month)+" del año "+str(year)


def getfiles(request):

    # Files (local path) to put in the .zip
    filenames=[]
    id_app = request.GET['id']
    app = Application.objects.get(pk = id_app)
    filenames.append(solicitud_facultad_doc(app))
    documents = app.get_documents()

    solicitante=str(app.id_Teacher)
    replacements=app.get_replacements()
    replacement_teachers=set()

    for replacement in replacements:
        replacement_teachers.add(replacement.rut_teacher)
    if len(replacement_teachers)==1:
        filenames.append(carta_reemplazo_ac_doc(app,replacement.rut_teacher))
    else:
        for replacement in replacements:
            filenames.append(carta_reemplazo_any(app,replacement))

    filenames.append(peticion_docente_doc(app,replacement_teachers))

    # Folder name in ZIP archive which contains the above files
    # E.g [thearchive.zip]/somefiles/file2.txt
    zip_subdir = "documentos"
    zip_filename = "%s.zip" % zip_subdir

    # Open StringIO to grab in-memory ZIP contents
    f = io.BytesIO()
    # The zip compressor
    zf = zipfile.ZipFile(f, "w")

    for fpath in filenames:
        # Calculate path for file in zip
        fdir, fname = os.path.split(fpath)
        zip_path = os.path.join(zip_subdir, fname)

        # Add file, at correct path
        zf.write(fpath, zip_path)

    for fpath in documents:
        # Calculate path for file in zip
        fdir, fname = os.path.split(fpath.file.name)
        zip_path = os.path.join(zip_subdir, fname)
        file_path = os.path.join(settings.MEDIA_ROOT,fpath.file.name)
        # Add file, at correct path
        zf.write(file_path, zip_path)
    # Must close zip for all contents to be written
    zf.close()

    # Grab ZIP file from in-memory, make response with correct MIME-type
    resp = HttpResponse(f.getvalue(),content_type='application/zip')
    # ..and correct content-disposition
    resp['Content-Disposition'] = 'attachment; filename=%s' % zip_filename

    return resp

def peticion_docente_doc(app,replacement_teachers):
    replacements=app.get_replacements()
    path = os.path.join(settings.MEDIA_ROOT, 'carta_peticion_docente.docx') #todo:path para cada profe (?)
    tipo_comision=str(app.id_commission_type)
    if tipo_comision == 'Academica':
        tipo_comision = 'Académica'
    con_o_sin_remuneracion="con" #TODO cuando es con? cuando es sin?
    destinos = app.get_destinations()
    fecha_inicio=datetimeToString(app.get_start_date())
    print(datetimeToString(app.get_start_date()))
    fecha_fin=datetimeToString(app.get_end_date())
    director = str(app.directors_name)
    nombre_solicitante = str(app.id_Teacher.name) +" "+ str(app.id_Teacher.last_name)
    email_solicitante = str(app.id_Teacher.mail)
    director = ""
    try:
        systemInformation = SystemInformation.objects.get(pk=1)
        director = str(systemInformation.director)
    except:
        pass

    intro="Por la presente ruego a Ud. tenga a bien autorizar una Comisión "+tipo_comision \
          +" "+ con_o_sin_remuneracion+" goce de remuneraciones, del "\
          +fecha_inicio+" al "+fecha_fin+". "
    if len(destinos) > 1:
        intro = intro + "Los motivos de esta solicitud son viajar a: "
        for destination in destinos:
            intro = intro + "A "+str(destination.city)+", "+str(destination.country)+" con objeto de "+ str(destination.motive)+". "
    else:
        intro = intro + "El motivo de esta solicitud es que deberé viajar a "+str(destinos[0].city)+", "+str(destinos[0].country)
        intro = intro + " es " + str(destinos[0].motive)+". "
    intro = intro+"Se adjuntan los documentos pertinentes."

    if len(replacement_teachers)==1:
        replacement_type="En mis actividades académicas y docentes seré reemplazado por "+str(replacement_teachers.pop())+"."
    else:
        i=0
        replace=[]
        for rep in replacements:
            if(str(rep.type)=="Docente"):
                replace.append("docentes")
            else:
                replace.append("académicas")
            i+=1
            replace.append(str(rep.rut_teacher))
        replacement_type="En mis actividades " +replace[0]+ " seré reemplazado por " +replace[1]+ " y en mis actividades "+ replace[2]+ " seré reemplazado por " +replace[3]+"."

    try:
        signature_file = app.id_Teacher.signature
        signature_path = str(signature_file)
        # print('firma')
        #print(signature_path)
        #signature_path = signature_path.split("/")
        #print(signature_path)
        #signature_path = signature_path[1]
        #print(signature_path)
    except:
        print('no hay firma')

    # user = request.user.id
    document = Document()
    document.add_paragraph(nombre_solicitante).alignment = 2 #
    document.add_paragraph("DCC - UChile").alignment = 2
    document.add_paragraph(email_solicitante).alignment = 2
    p = document.add_paragraph()
    p.add_run().add_break()
    p.add_run('Señor ').bold = True
    p.add_run().add_break()
    p.add_run(director).bold = True
    p.add_run().add_break()
    p.add_run('Director DCC').bold = True
    p.add_run().add_break()
    p.add_run('Presente').bold = True
    p.add_run().add_break()

    intro=document.add_paragraph(intro)
    intro.add_run().add_break()
    finances= app.get_finances()
    prefix_financiamiento = "El financiamiento consistirá: "
    financiamiento = ""
    for finance in finances:
        if finance.financed_by != None:
            finance_type = str(finance.id_finance_type.type)
            if finance_type == 'Viatico':
                finance_type = 'Viático'
            if finance_type == 'Inscripcion':
                finance_type = 'Inscripción'
            financiamiento = financiamiento+"Los gastos de "+finance_type
            if finance.amount != None:
                financiamiento = financiamiento + " de $"+str(finance.amount)+" "+str(finance.id_currency.currency)+" "
            financiamiento = financiamiento + "serán financiados por " + str(finance.financed_by)+". "
    if len(financiamiento) > 0:
        financiamiento=document.add_paragraph(prefix_financiamiento+financiamiento)
        financiamiento.add_run().add_break()
    reemplazo=document.add_paragraph(replacement_type)
    saludo = "Sin otro particular, le saluda atentamente,"
    saludo=document.add_paragraph(saludo)
    saludo.add_run().add_break()
    saludo.add_run().add_break()
    saludo.add_run().add_break()
    saludo.add_run().add_break()
    espacio='                                                                                                                                                  '
    saludo.add_run(espacio)
    try:
        saludo.add_run().add_picture(os.path.join(settings.MEDIA_ROOT,signature_path.split('/')[0],signature_path.split('/')[1]),width=Inches(1.0)).alignment = WD_ALIGN_PARAGRAPH.RIGHT#signature_file.path)
    except:
        print("error al agregar la firma")
    saludo.add_run().add_break()
    saludo.add_run().add_break()
    #profe=str(app.id_Teacher)
    profe=document.add_paragraph(nombre_solicitante)
    profe.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    document.save(path)
    return path

def solicitud_facultad_doc(app):
    teacher=app.id_Teacher
    document = Document()
    pretitle=document.add_paragraph()
    pretitle.add_run("OFICIO N°")
    pretitle.add_run().add_break()
    pretitle.add_run("FECHA:   "+datetimeToString(date.today()))
    pretitle.add_run().add_break()
    pretitle.add_run("USO INTERNO FACULTAD").bold = True
    pretitle.add_run().add_break()
    pretitle.alignment = WD_ALIGN_PARAGRAPH.RIGHT


    title = document.add_paragraph('SOLICITUD DE COMISION O PERMISO')
    title.alignment = 1
    p = document.add_paragraph()
    run = p.add_run()
    run.add_break()
    p.add_run('UNIDAD ACADEMICA         ').bold = True
    p.add_run('DEPARTAMENTO DE CIENCIAS DE LA COMPUTACION')
    p.add_run().add_break()
    p.add_run('NOMBRE                                  ').bold = True
    p.add_run(str(teacher).upper())
    p.add_run('            RUT   ').bold = True
    p.add_run(str(teacher.rut))
    p.add_run().add_break()
    p.add_run('CARGO                                      ').bold = True
    p.add_run('ACADEMICO JORNADA '+str(teacher.working_day).upper())#TODO jornada?
    p.add_run('          GRADO ').bold = True
    p.add_run().add_break()
    p.add_run('JERARQUIA                            ').bold = True
    p.add_run(str(teacher.hierarchy).upper())
    p.add_run().add_break()
    p.add_run('TIPO DE MOVIMIENTO     ').bold = True
    p.add_run('COMISION '+str(app.id_commission_type).upper())#TODO remuneraciones
    p.add_run().add_break()
    for destination in app.get_destinations():
        p.add_run().add_break()
        p.add_run('PERIODO                                ').bold = True
        p.add_run('DEL     ').bold = True
        p.add_run(datetimeToString(destination.start_date))
        p.add_run('    AL    ').bold = True
        p.add_run(datetimeToString(destination.end_date))
        p.add_run().add_break()
        p.add_run('LUGAR                                     ').bold = True
        lugar = str(destination.city)+", "+str(destination.country)
        p.add_run(lugar.upper())
        p.add_run().add_break()
        p.add_run().add_break()
        objectivo = str(destination.motive)
        p.add_run('OBJETIVO                              ').bold = True
        p.add_run(objectivo.upper())
        #p.add_run(str(app.motive))
        p.add_run().add_break()
    p.add_run().add_break()
    p.add_run('FUENTES DE FINANCIAMIENTO (INDICAR MONTO Y ORIGEN)').bold = True
    p.add_run().add_break()


    p.add_run('INCRIPCIÓN                         ').bold = True
    try:
        finance=Finance.objects.get(id_application=app,id_finance_type=1)
        if finance.id_currency!=None and finance.amount!=None and finance.financed_by!=None:
            financiamiento = str(finance.amount)+" "+str(finance.id_currency.currency)+", "+str(finance.financed_by)
            p.add_run(financiamiento.upper())
        else:
            p.add_run(str("NO SOLICITA"))
    except:
        p.add_run(str("NO SOLICITA"))
    p.add_run().add_break()
    p.add_run('PASAJE IDA Y VUELTA    ').bold = True
    try:
        finance=Finance.objects.get(id_application=app,id_finance_type=2)
        if finance.id_currency!=None and finance.amount!=None and finance.financed_by!=None:
            financiamiento = str(finance.amount)+" "+str(finance.id_currency.currency)+", "+str(finance.financed_by)
            p.add_run(financiamiento.upper())
        else:
            p.add_run(str("NO SOLICITA"))
    except:
        p.add_run(str("NO SOLICITA"))
    p.add_run().add_break()
    p.add_run('AYUDA DE VIAJE                ').bold = True
    try:
        finance=Finance.objects.get(id_application=app,id_finance_type=3)
        if finance.id_currency!=None and finance.amount!=None and finance.financed_by!=None:
            financiamiento = str(finance.amount)+" "+str(finance.id_currency.currency)+", "+str(finance.financed_by)
            p.add_run(financiamiento.upper())
        else:
            p.add_run(str("NO SOLICITA"))
    except:
        p.add_run(str("NO SOLICITA"))
    p.add_run().add_break()
    p.add_run().add_break()
    p.add_run('OBSERVACIONES ').bold = True
    p.add_run().add_break()
    p.add_run().add_break()
    p.add_run('REEMPLAZANTE DE CATEDRAS     ').bold = True
    p.add_run(str(Replacement.objects.get(id_Application=app,type=1).rut_teacher).upper())
    p.add_run().add_break()
    p.add_run('REEMPLAZANTE DE CARGO             ').bold = True
    p.add_run(str(Replacement.objects.get(id_Application=app, type=2).rut_teacher).upper())
    p.add_run().add_break()
    p.add_run().add_break()
    p.add_run().add_break()
    p.add_run().add_break()
    p.add_run().add_break()
    p.add_run().add_break()
    firmadir=document.add_paragraph()
    director = ""
    try:
        systemInformation = SystemInformation.objects.get(pk=1)
        director = str(systemInformation.director)
    except:
        pass
    firmadir.add_run('             '+director.upper()+'                ').bold = True
    firmadir.add_run().add_break()
    firmadir.add_run('                DIRECTOR                  ')
    firmadir.add_run().add_break()
    firmadir.add_run('DEPARTAMENTO DE CIENCIAS DE LA COMPUTACIÓN')
    firmadir.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    path = os.path.join(settings.MEDIA_ROOT, 'solicitud_facultad.doc')
    document.save(path)
    return path

def carta_reemplazo_any(app,replacement):
    solicitante=str(app.id_Teacher)
    replacement_teacher=replacement.rut_teacher
    replacement_type=replacement.type
    return carta_reemplazo(solicitante,app,replacement_teacher,replacement_type)

def carta_reemplazo_ac_doc(app,replacement_teacher):
    solicitante=str(app.id_Teacher)
    replacement_type="académicas y docentes"

    return carta_reemplazo(solicitante,app,replacement_teacher,replacement_type)

def carta_reemplazo(solicitante,app,replacement_teacher,replacement_type):

    signature_file=""
    try:
        signature_file=replacement_teacher.signature
        signature_path=str(signature_file)
        #print('firma')
        #print(signature_path)
        #signature_path=signature_path.split("/")
        #print(signature_path)
        #signature_path=signature_path[1]
        #print(signature_path)
    except:
        print('error obteniendo la firma')
        print(traceback.format_exc())
    fecha_inicio=datetimeToString(app.get_start_date())
    fecha_fin=datetimeToString(app.get_end_date())

    filename= 'compromiso_reemplazo'+str(replacement_type)+'.doc'
    path = os.path.join(settings.MEDIA_ROOT,filename)
    document=Document()

    today="Santiago, "+datetimeToString(date.today())
    today_date=document.add_paragraph(today)
    today_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    today_date.add_run().add_break()
    tipo_reemplazo = str(replacement_type).upper()
    if tipo_reemplazo == 'DOCENTE':
        tipo_reemplazo = 'DOCENTES'
    if tipo_reemplazo == 'ACADEMICO':
        tipo_reemplazo = 'ACADÉMICAS'
    title = document.add_paragraph('CARTA DE COMPROMISO DE REEMPLAZO EN FUNCIONES '+tipo_reemplazo)
    title.alignment = 1
    title.add_run().add_break()
    title.add_run().add_break()
    title.add_run().add_break()
    nombre_reemplazante = str(replacement_teacher.name)+" "+str(replacement_teacher.last_name)
    cuerpo='Yo, '+nombre_reemplazante+', me comprometo a reemplazar a '+solicitante+', del '+fecha_inicio+ ' al '+fecha_fin+', ambas fechas inclusive, en todas sus actividades del tipo '+str(replacement_type)+'.'
    #jerarquia=replacement_teacher.hierarchy
    #jornada=replacement_teacher.working_day

    paragraph = document.add_paragraph(cuerpo)
    paragraph.add_run().add_break()
    paragraph.add_run().add_break()
    paragraph.add_run().add_break()

    p = document.add_paragraph()
    p.add_run().add_break()
    p.add_run('Nombre                           : ').bold = True
    p.add_run(nombre_reemplazante)
    p.add_run().add_break()
    p.add_run('Jerarquía                        : ').bold = True
    p.add_run(str(replacement_teacher.hierarchy))
    p.add_run().add_break()
    p.add_run('Cargo                                : ').bold = True
    p.add_run('Académico Jornada ' +str(replacement_teacher.get_working_day()))
    #p.add_run(str(jornada))
    p.add_run().add_break()
    p.add_run().add_break()
    p.add_run().add_break()
    p.add_run().add_break()
    p.add_run('                                                                                                               ')

    try:
        p.add_run().add_picture(os.path.join(settings.MEDIA_ROOT,signature_path.split('/')[0],signature_path.split('/')[1]),width=Inches(1.0)).alignment = WD_ALIGN_PARAGRAPH.RIGHT#signature_file.path)
    except:
        print("error al agregar la firma")

    signature=document.add_paragraph()
    signature.add_run('                                                                  '+nombre_reemplazante).bold=True
    signature.alignment = WD_ALIGN_PARAGRAPH.CENTER
    signature.add_run().add_break()
    signature.add_run().add_break()
    signature.add_run().add_break()

    other_signatures=document.add_paragraph()
    other_signatures.add_run('V°B° Sr. Director').bold = True
    other_signatures.add_run().add_break()
    other_signatures.add_run().add_break()
    other_signatures.add_run('V°B° Jefa Docente').bold = True
    #document.add_paragraph('first item in unordered list', style='ListBullet')
    #document.add_heading('The role of dolphins', level=2)
    #document.add_picture('image-filename.png')
    document.save(path)
    return path


def getattachedfiles(request):

    id_app = request.GET['id']
    app = Application.objects.get(pk = id_app)
    documents = app.get_documents()

    # Folder name in ZIP archive which contains the above files
    # E.g [thearchive.zip]/somefiles/file2.txt
    zip_subdir = "documentos"
    zip_filename = "%s.zip" % zip_subdir

    # Open StringIO to grab in-memory ZIP contents
    f = io.BytesIO()
    # The zip compressor
    zf = zipfile.ZipFile(f, "w")

    for fpath in documents:
        # Calculate path for file in zip
        fdir, fname = os.path.split(fpath.file.name)
        zip_path = os.path.join(zip_subdir, fname)
        file_path = os.path.join(settings.MEDIA_ROOT,fpath.file.name)
        # Add file, at correct path
        zf.write(file_path, zip_path)
    # Must close zip for all contents to be written
    zf.close()

    # Grab ZIP file from in-memory, make response with correct MIME-type
    resp = HttpResponse(f.getvalue(),content_type='application/zip')
    # ..and correct content-disposition
    resp['Content-Disposition'] = 'attachment; filename=%s' % zip_filename

    return resp